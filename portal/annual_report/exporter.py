"""Export annual report analysis to DOCX, PDF, and XLSX."""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Any, Dict, List

import fitz

from .metrics_extractor import metrics_to_rows


def export_docx(markdown: str, title: str = "年报财报分析") -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as e:
        raise ImportError("请安装 python-docx: pip install python-docx") from e

    doc = Document()
    doc.add_heading(title, 0)

    for line in markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), 1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), 2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), 3)
        elif stripped.startswith("|") and stripped.endswith("|"):
            # skip table rendering in simple mode — add as paragraph
            p = doc.add_paragraph(stripped)
            p.runs[0].font.size = Pt(9)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif stripped.startswith(">"):
            doc.add_paragraph(stripped.lstrip("> ").strip())
        else:
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_pdf(markdown: str, title: str = "年报财报分析") -> bytes:
    """Plain-text PDF via PyMuPDF."""
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    margin = 50
    rect = fitz.Rect(margin, margin, 595 - margin, 842 - margin)
    plain = re.sub(r"\*\*(.+?)\*\*", r"\1", markdown)
    plain = re.sub(r"^#+\s*", "", plain, flags=re.M)
    text = f"{title}\n\n{plain}"
    fontname = "china-s"
    try:
        overflow = page.insert_textbox(rect, text, fontsize=9, fontname=fontname, align=fitz.TEXT_ALIGN_LEFT)
    except Exception:
        fontname = "helv"
        overflow = page.insert_textbox(rect, text, fontsize=9, fontname=fontname, align=fitz.TEXT_ALIGN_LEFT)
    while overflow:
        page = doc.new_page(width=595, height=842)
        rect = fitz.Rect(margin, margin, 595 - margin, 842 - margin)
        overflow = page.insert_textbox(rect, overflow, fontsize=9, fontname=fontname)
    out = doc.tobytes()
    doc.close()
    return out


def export_xlsx(metrics: Dict[str, Any]) -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "关键指标"
    for row in metrics_to_rows(metrics):
        ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def export_filename(base: str, ext: str) -> str:
    safe = re.sub(r'[<>:"/\\|?*]', "_", base)[:80]
    return f"{safe}.{ext}"
