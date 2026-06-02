"""Extract plain text from tender attachment files (PDF / Word)."""

from __future__ import annotations

from pathlib import Path

MAX_ATTACHMENT_CHARS = 200_000


def extract_text_from_file(path: Path) -> str:
    """Return UTF-8 text from PDF or Word; empty string if unsupported."""
    if not path.is_file():
        return ""

    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return _extract_pdf(path)
        if suffix in (".docx", ".doc"):
            return _extract_docx(path)
    except Exception:
        return ""
    return ""


def _extract_pdf(path: Path) -> str:
    import fitz

    parts: list[str] = []
    with fitz.open(path) as doc:
        for page in doc:
            parts.append(page.get_text() or "")
    text = "\n".join(parts).strip()
    return text[:MAX_ATTACHMENT_CHARS]


def _extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    text = "\n".join(parts).strip()
    return text[:MAX_ATTACHMENT_CHARS]
